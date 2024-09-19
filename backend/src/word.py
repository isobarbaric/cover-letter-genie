from pathlib import Path
from typing import Literal, List, Dict, Tuple, get_args
from docx import Document
from utils.marker_utils import clean_markers, determine_answer
import re

WordFileExtension = Literal["docx"]
SUPPORTED_WORD_FILE_EXTENSIONS = get_args(WordFileExtension)

def get_markers(docx_path: Path) -> Tuple[Dict[str, str], List[str]]:
    # starting this indexing at 1 to shave off the '.' to get the actual extension
    filename_suffix = docx_path.suffix.lower()[1:] 
    if filename_suffix not in SUPPORTED_WORD_FILE_EXTENSIONS:
        raise ValueError(f"Unsupported file type {filename_suffix}; only .docx files are supported")

    new_to_old_marker = {}

    doc = Document(docx_path)
    context = []
    for para in doc.paragraphs:
        text = para.text
        markers = set(re.findall(r"\[.*?\]", text))
        local = [word[1:-1] for word in markers]
        # for marker in local:
        #     if marker not in new_to_old_marker:
        #         new_to_old_marker[marker] = marker
        context += local

    cleaned_markers = clean_markers(context)
    for idx in range(len(cleaned_markers)):
        old_marker = context[idx]
        new_marker = cleaned_markers[idx]
        new_to_old_marker[new_marker] = old_marker

    return new_to_old_marker, cleaned_markers


def replace_markers(
    pdf_file_path: Path,
    answers: Dict[str, str],
    markers: list[str],
    marker_map: Dict[str, str],
    convert_to_pdf: bool = True
):
    # determine markers' replacements
    replacements = []
    for marker in markers:
        replacements.append((marker, answers[marker]))

    # now perform the replacements
    doc = Document(pdf_file_path)
    for marker, answer in replacements:
        marker = "[" + marker_map[marker] + "]"
        doc = Document(pdf_file_path)
        for para in doc.paragraphs:
            for run in para.runs:
                for marker, answer in replacements:
                    full_marker = f"[{marker_map[marker]}]"
                    if full_marker in run.text:
                        run.text = run.text.replace(full_marker, answer)

    # TODO: add an output directory/output file parameter 
    output_path = Path("output") / "test.docx"
    if not output_path.parent.exists():
        output_path.parent.mkdir(parents=True)

    # print(doc_path)
    doc.save(output_path)

    if convert_to_pdf:
        output_path = Path("output") / "test2.pdf"
        # save_as_pdf(doc_path, output_path)

# TODO: add a working function to convert the docx file to pdf
def save_as_pdf(
    doc_path: Path, 
    output_path: Path
):
    pass
    # convert(str(output_path))
    # word = comtypes.client.CreateObject('Word.Application')
    # word.Visible = False
    # doc = word.Documents.Open(str(doc_path))
    # doc.SaveAs(str(output_path), FileFormat=17)
    # doc.Close()
    # word.Quit()

    # doc_path.unlink()
        

if __name__ == "__main__":
    pdf_path = Path("templates") / "template-2.docx"

    # eventually, will make this a prompt(?)
    wworks_desc = Path("wworks_page.txt")

    # marker_map, markers = get_markers(pdf_path)
    # print(markers, marker_map)

    answer = determine_answer(wworks_desc, markers)
    # replace_markers(pdf_path, answer, markers, marker_map)